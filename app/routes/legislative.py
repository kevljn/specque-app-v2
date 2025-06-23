from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, make_response, abort
from flask_login import login_required, current_user
from app.models.legislative import LegislativeText, Amendment, Vote, LegalElement
from app import db
from app.forms import LegislativeTextForm
import time
import io
from weasyprint import HTML
from docx import Document
from odf.opendocument import OpenDocumentText
from odf.text import P, Span, H
from odf.style import Style, TextProperties
from flask_wtf.csrf import CSRFProtect
import difflib
from markupsafe import Markup

legislative_bp = Blueprint('legislative', __name__)
csrf = CSRFProtect()

def html_diff(a, b):
    diff = difflib.ndiff(a.split(), b.split())
    result = []
    for word in diff:
        if word.startswith('- '):
            result.append(f'<strong><del>{word[2:]}</del></strong>')
        elif word.startswith('+ '):
            continue  # On n'affiche pas les ajouts côté texte d'origine
        elif word.startswith('  '):
            result.append(word[2:])
    return Markup(' '.join(result))

def html_diff_deletions(a, b):
    diff = difflib.ndiff(a.split(), b.split())
    result = []
    for word in diff:
        if word.startswith('- '):
            result.append(f'<strong><del>{word[2:]}</del></strong>')
        elif word.startswith('+ '):
            continue
        elif word.startswith('  '):
            result.append(word[2:])
    return Markup(' '.join(result))

def html_diff_additions(a, b):
    diff = difflib.ndiff(a.split(), b.split())
    result = []
    for word in diff:
        if word.startswith('+ '):
            result.append(f'<strong>{word[2:]}</strong>')
        elif word.startswith('- '):
            continue
        elif word.startswith('  '):
            result.append(word[2:])
    return Markup(' '.join(result))

@legislative_bp.route('/texts/new', methods=['GET', 'POST'])
@login_required
def new_text():
    if not current_user.is_admin() and not current_user.is_reporter():
        flash('Accès refusé : seuls les admins ou rapporteurs peuvent créer un texte.', 'error')
        return redirect(url_for('main.dashboard'))
    form = LegislativeTextForm()
    if form.validate_on_submit():
        text = LegislativeText(
            title=form.title.data,
            content=form.content.data,
            status='draft'
        )
        db.session.add(text)
        db.session.commit()
        flash('Texte législatif créé avec succès.', 'success')
        return redirect(url_for('legislative.view_text', text_id=text.id))
    return render_template('legislative/new_text.html', form=form)

@legislative_bp.route('/texts/<int:text_id>')
@login_required
def view_text(text_id):
    text = LegislativeText.query.get_or_404(text_id)
    # Charger la hiérarchie LegalElement
    elements = LegalElement.query.filter_by(legislative_text_id=text_id, parent_id=None).all()
    # Charger tous les amendements liés à ce texte
    amendments = Amendment.query.filter_by(legislative_text_id=text_id).all()
    amendments_by_element = {}
    for amend in amendments:
        if amend.legal_element_id:
            amendments_by_element.setdefault(amend.legal_element_id, []).append(amend)
    # Appliquer le dernier amendement (edit) à chaque LegalElement
    def apply_amendments(element):
        amends = amendments_by_element.get(element.id, [])
        edit_amends = [a for a in amends if a.type == 'edit']
        if edit_amends:
            # On prend le dernier
            element.amended_content = edit_amends[-1].content
        else:
            element.amended_content = None
        for child in element.children:
            apply_amendments(child)
    for el in elements:
        apply_amendments(el)
    return render_template('legislative/view_text.html', text=text, elements=elements, amendments_by_element=amendments_by_element, html_diff_deletions=html_diff_deletions, html_diff_additions=html_diff_additions)

@legislative_bp.route('/texts/<int:text_id>/amendments/new', methods=['GET', 'POST'])
@login_required
def new_amendment(text_id):
    text = LegislativeText.query.get_or_404(text_id)
    if request.method == 'POST':
        if request.is_json:
            data = request.get_json()
            content = data.get('modified_text')
            amendment_type = data.get('type', 'edit')
            legal_element_id = data.get('element_id')
        else:
            content = request.form.get('content')
            amendment_type = 'edit'
            legal_element_id = request.form.get('element_id')

        if not content:
            return jsonify({'status': 'error', 'message': 'Le contenu de l\'amendement ne peut pas être vide'}), 400

        amendment = Amendment(
            legislative_text_id=text_id,
            author_id=current_user.id,
            content=content,
            type=amendment_type,
            legal_element_id=legal_element_id
        )
        db.session.add(amendment)
        db.session.commit()
        
        if request.is_json:
            return jsonify({'status': 'success', 'message': 'Amendement proposé avec succès', 'amendment_id': amendment.id})
        else:
            flash('Amendement proposé avec succès.', 'success')
            return redirect(url_for('legislative.view_text', text_id=text_id))
            
    return render_template('legislative/new_amendment.html', text=text)

@legislative_bp.route('/amendments/<int:amendment_id>/vote', methods=['POST'])
@login_required
def vote_amendment(amendment_id):
    amendment = Amendment.query.get_or_404(amendment_id)
    vote_type = request.form.get('vote_type')
    
    # Vérifier si l'utilisateur a déjà voté
    existing_vote = Vote.query.filter_by(
        user_id=current_user.id,
        amendment_id=amendment_id
    ).first()
    
    if existing_vote:
        existing_vote.vote_type = vote_type
    else:
        vote = Vote(
            user_id=current_user.id,
            legislative_text_id=amendment.legislative_text_id,
            amendment_id=amendment_id,
            vote_type=vote_type
        )
        db.session.add(vote)
    
    db.session.commit()
    return jsonify({'status': 'success'})

@legislative_bp.route('/texts/<int:text_id>/status', methods=['POST'])
@login_required
def update_text_status(text_id):
    if not current_user.is_admin() and not current_user.is_reporter():
        flash('Accès non autorisé.', 'error')
        return redirect(url_for('main.dashboard'))
    
    text = LegislativeText.query.get_or_404(text_id)
    new_status = request.form.get('status')
    
    if new_status in ['draft', 'in_commission', 'in_plenary', 'adopted', 'rejected']:
        text.status = new_status
        db.session.commit()
        flash('Statut mis à jour avec succès.', 'success')
    
    return redirect(url_for('legislative.view_text', text_id=text_id))

@legislative_bp.route('/texts/<int:text_id>/article/<int:idx>/amend', methods=['POST'])
@login_required
def amend_article(text_id, idx):
    # Anti-spam : 1 action par 10s par utilisateur sur cet article
    last_action = getattr(current_user, 'last_action', None)
    now = time.time()
    if last_action and now - last_action < 10:
        return jsonify({'status': 'error', 'message': 'Trop d\'actions, veuillez patienter.'}), 429
    current_user.last_action = now
    content = request.form.get('content')
    amendment = Amendment(
        legislative_text_id=text_id,
        author_id=current_user.id,
        content=content,
        paragraph_index=idx,
        type='edit'
    )
    db.session.add(amendment)
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Amendement proposé.', 'amendment_id': amendment.id})

@legislative_bp.route('/texts/<int:text_id>/article/<int:idx>/add', methods=['POST'])
@login_required
def add_article(text_id, idx):
    last_action = getattr(current_user, 'last_action', None)
    now = time.time()
    if last_action and now - last_action < 10:
        return jsonify({'status': 'error', 'message': 'Trop d\'actions, veuillez patienter.'}), 429
    current_user.last_action = now
    new_content = request.form.get('content')
    text = LegislativeText.query.get_or_404(text_id)
    amendment = Amendment(
        legislative_text_id=text_id,
        author_id=current_user.id,
        content=new_content,
        paragraph_index=idx+1,
        type='add'
    )
    db.session.add(amendment)
    db.session.commit()
    return jsonify({
        'status': 'success',
        'message': 'Article ajouté.',
        'amendment_id': amendment.id,
        'insert_idx': idx+1,
        'article_text': new_content
    })

@legislative_bp.route('/texts/<int:text_id>/article/<int:idx>/delete', methods=['POST'])
@login_required
def delete_article(text_id, idx):
    last_action = getattr(current_user, 'last_action', None)
    now = time.time()
    if last_action and now - last_action < 10:
        return jsonify({'status': 'error', 'message': 'Trop d\'actions, veuillez patienter.'}), 429
    current_user.last_action = now
    amendment = Amendment(
        legislative_text_id=text_id,
        author_id=current_user.id,
        content='',
        paragraph_index=idx,
        type='delete'
    )
    db.session.add(amendment)
    db.session.commit()
    return jsonify({'status': 'success', 'message': 'Amendement de suppression proposé.', 'amendment_id': amendment.id})

@legislative_bp.route('/texts/<int:text_id>/download/pdf')
@login_required
def download_pdf(text_id):
    text = LegislativeText.query.get_or_404(text_id)
    user_amendments = { (a.paragraph_index, a.type): a for a in Amendment.query.filter_by(legislative_text_id=text_id, author_id=current_user.id).all() }
    articles = [a for a in text.content.split('\n') if a.strip()]
    display_articles = []
    for idx, article in enumerate(articles):
        if (idx, 'delete') in user_amendments:
            display_articles.append({'text': article, 'deleted': True, 'original': article, 'modified': None})
        elif (idx, 'edit') in user_amendments:
            display_articles.append({'text': user_amendments[(idx, 'edit')].content, 'deleted': False, 'original': article, 'modified': user_amendments[(idx, 'edit')].content})
        else:
            display_articles.append({'text': article, 'deleted': False, 'original': article, 'modified': None})
    add_amendments = sorted([a for a in user_amendments.values() if a.type == 'add'], key=lambda a: a.paragraph_index if a.paragraph_index is not None else 9999)
    for add in reversed(add_amendments):
        insert_idx = add.paragraph_index + 1 if add.paragraph_index is not None else len(display_articles)
        display_articles.insert(insert_idx, {'text': add.content, 'deleted': False, 'original': None, 'modified': add.content})
    html = f'<h1>{text.title}</h1>'
    for idx, art in enumerate(display_articles, 1):
        if art['deleted']:
            html += f'<p style="text-decoration:line-through;color:#888;"><b>Article {idx}.</b> {art["text"]}</p>'
        elif art['modified']:
            html += f'<p><b>Article {idx}.</b> <i>{art["original"]}</i><br><span>{art["modified"]}</span></p>'
        else:
            html += f'<p><b>Article {idx}.</b> {art["text"]}</p>'
    pdf_io = io.BytesIO()
    HTML(string=html).write_pdf(pdf_io)
    pdf_io.seek(0)
    return send_file(pdf_io, mimetype='application/pdf', as_attachment=True, download_name=f'{text.title}.pdf')

@legislative_bp.route('/texts/<int:text_id>/download/word')
@login_required
def download_word(text_id):
    text = LegislativeText.query.get_or_404(text_id)
    user_amendments = { (a.paragraph_index, a.type): a for a in Amendment.query.filter_by(legislative_text_id=text_id, author_id=current_user.id).all() }
    articles = [a for a in text.content.split('\n') if a.strip()]
    display_articles = []
    for idx, article in enumerate(articles):
        if (idx, 'delete') in user_amendments:
            display_articles.append({'text': article, 'deleted': True, 'original': article, 'modified': None})
        elif (idx, 'edit') in user_amendments:
            display_articles.append({'text': user_amendments[(idx, 'edit')].content, 'deleted': False, 'original': article, 'modified': user_amendments[(idx, 'edit')].content})
        else:
            display_articles.append({'text': article, 'deleted': False, 'original': article, 'modified': None})
    add_amendments = sorted([a for a in user_amendments.values() if a.type == 'add'], key=lambda a: a.paragraph_index if a.paragraph_index is not None else 9999)
    for add in reversed(add_amendments):
        insert_idx = add.paragraph_index + 1 if add.paragraph_index is not None else len(display_articles)
        display_articles.insert(insert_idx, {'text': add.content, 'deleted': False, 'original': None, 'modified': add.content})
    doc = Document()
    doc.add_heading(text.title, 0)
    for idx, art in enumerate(display_articles, 1):
        if art['deleted']:
            p = doc.add_paragraph()
            p.add_run(f'Article {idx}. {art["text"]}').font.strike = True
        elif art['modified']:
            p = doc.add_paragraph()
            p.add_run(f'Article {idx}. ').bold = True
            p.add_run(art['original']).italic = True
            p.add_run(' -> ')
            p.add_run(art['modified'])
        else:
            doc.add_paragraph(f'Article {idx}. {art["text"]}')
    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return send_file(word_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f'{text.title}.docx')

@legislative_bp.route('/texts/<int:text_id>/download/odt')
@login_required
def download_odt(text_id):
    text = LegislativeText.query.get_or_404(text_id)
    user_amendments = { (a.paragraph_index, a.type): a for a in Amendment.query.filter_by(legislative_text_id=text_id, author_id=current_user.id).all() }
    articles = [a for a in text.content.split('\n') if a.strip()]
    display_articles = []
    for idx, article in enumerate(articles):
        if (idx, 'delete') in user_amendments:
            display_articles.append({'text': article, 'deleted': True, 'original': article, 'modified': None})
        elif (idx, 'edit') in user_amendments:
            display_articles.append({'text': user_amendments[(idx, 'edit')].content, 'deleted': False, 'original': article, 'modified': user_amendments[(idx, 'edit')].content})
        else:
            display_articles.append({'text': article, 'deleted': False, 'original': article, 'modified': None})
    add_amendments = sorted([a for a in user_amendments.values() if a.type == 'add'], key=lambda a: a.paragraph_index if a.paragraph_index is not None else 9999)
    for add in reversed(add_amendments):
        insert_idx = add.paragraph_index + 1 if add.paragraph_index is not None else len(display_articles)
        display_articles.insert(insert_idx, {'text': add.content, 'deleted': False, 'original': None, 'modified': add.content})
    
    # Créer le document ODT
    odt = OpenDocumentText()
    
    # Ajouter le titre
    title_style = Style(name="Title", family="paragraph")
    title_style.addElement(TextProperties(attributes={'fontsize': "24pt", 'fontweight': "bold"}))
    odt.automaticstyles.addElement(title_style)
    
    title_para = P(stylename=title_style)
    title_para.addElement(H(text=text.title))
    odt.text.addElement(title_para)
    
    # Style pour les articles normaux
    normal_style = Style(name="Normal", family="paragraph")
    odt.automaticstyles.addElement(normal_style)
    
    # Style pour les articles supprimés
    deleted_style = Style(name="Deleted", family="paragraph")
    deleted_style.addElement(TextProperties(attributes={'textlinethroughstyle': "solid", 'textlinethroughcolor': "#888888"}))
    odt.automaticstyles.addElement(deleted_style)
    
    # Style pour les articles modifiés
    modified_style = Style(name="Modified", family="paragraph")
    modified_style.addElement(TextProperties(attributes={'fontstyle': "italic", 'color': "#888888"}))
    odt.automaticstyles.addElement(modified_style)
    
    # Ajouter les articles
    for idx, art in enumerate(display_articles, 1):
        if art['deleted']:
            p = P(stylename=deleted_style)
            p.addElement(H(text=f"Article {idx}. {art['text']}"))
            odt.text.addElement(p)
        elif art['modified']:
            # Article original
            p = P(stylename=modified_style)
            p.addElement(H(text=f"Article {idx}. {art['original']}"))
            odt.text.addElement(p)
            # Article modifié
            p = P(stylename=normal_style)
            p.addElement(H(text=art['modified']))
            odt.text.addElement(p)
        else:
            p = P(stylename=normal_style)
            p.addElement(H(text=f"Article {idx}. {art['text']}"))
            odt.text.addElement(p)
    
    # Sauvegarder le document
    odt_io = io.BytesIO()
    odt.write(odt_io)
    odt_io.seek(0)
    
    return send_file(
        odt_io,
        mimetype='application/vnd.oasis.opendocument.text',
        as_attachment=True,
        download_name=f'{text.title}.odt'
    )

# Undo global (par id)
@legislative_bp.route('/amendments/<int:amendment_id>/undo', methods=['POST'])
@login_required
def undo_amendment(amendment_id):
    amendment = Amendment.query.get_or_404(amendment_id)
    if amendment.author_id != current_user.id:
        flash('Accès refusé : vous ne pouvez annuler que vos propres amendements.', 'error')
        abort(403)
    db.session.delete(amendment)
    db.session.commit()
    flash('Amendement annulé.', 'success')
    return '', 204

# Undo par article (supprime tous les amendements de l'utilisateur sur cet article)
@legislative_bp.route('/texts/<int:text_id>/article/<int:idx>/undo', methods=['POST'])
@login_required
def undo_article_amendment(text_id, idx):
    amendments = Amendment.query.filter_by(legislative_text_id=text_id, author_id=current_user.id, paragraph_index=idx).all()
    for a in amendments:
        db.session.delete(a)
    db.session.commit()
    return '', 204 